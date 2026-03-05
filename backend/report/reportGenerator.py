"""
reportGenerator.py – Generates .docx security assessment reports

Compiler & AI Based Bug-Detector

Uses python-docx to create professional, structured reports from
the canonical JSON analysis output.
"""

import json
import sys
import os
from datetime import datetime
from typing import Dict, List

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("[Report Generator] python-docx not installed. Install with: pip install python-docx",
          file=sys.stderr)


# Grade colors
GRADE_COLORS = {
    "Secure": RGBColor(16, 185, 129) if HAS_DOCX else None,        # Green
    "Moderate Risk": RGBColor(245, 158, 11) if HAS_DOCX else None,  # Amber
    "High Risk": RGBColor(239, 68, 68) if HAS_DOCX else None,       # Red
}

SEVERITY_COLORS = {
    "critical": RGBColor(220, 38, 38) if HAS_DOCX else None,
    "high": RGBColor(239, 68, 68) if HAS_DOCX else None,
    "medium": RGBColor(245, 158, 11) if HAS_DOCX else None,
    "low": RGBColor(59, 130, 246) if HAS_DOCX else None,
    "info": RGBColor(148, 163, 184) if HAS_DOCX else None,
}


def generate_report(report_data: Dict, output_path: str) -> str:
    """
    Generate a .docx report from canonical JSON data.

    Args:
        report_data: The canonical analysis JSON
        output_path: Path to save the .docx file

    Returns:
        Path to the generated .docx file
    """
    if not HAS_DOCX:
        raise ImportError("python-docx is required for report generation")

    doc = Document()

    # -- Style setup --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # == Section 1: Cover Page ==
    _add_cover_page(doc, report_data)

    # == Section 2: File Info & Code Statistics ==
    _add_file_info(doc, report_data)

    # == Section 3: Risk Score Analysis ==
    _add_risk_score(doc, report_data)

    # == Section 4: OWASP Summary Table ==
    _add_owasp_summary(doc, report_data.get("findings", []))

    # == Section 5: Detailed Findings ==
    _add_detailed_findings(doc, report_data.get("findings", []))

    # Save
    doc.save(output_path)
    return output_path


def _add_cover_page(doc: Document, data: Dict):
    """Add the cover page with title, timestamp, and risk score."""
    doc.add_paragraph("")
    doc.add_paragraph("")

    title = doc.add_heading("Compiler & AI Based Bug-Detector", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading("Security Assessment Report", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"File: {data.get('file', 'N/A')}").bold = True
    info.add_run(f"\nLanguage: {data.get('language', 'N/A').upper()}")
    info.add_run(f"\nJob ID: {data.get('job_id', 'N/A')}")
    info.add_run(f"\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_paragraph("")

    score_para = doc.add_paragraph()
    score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    score_run = score_para.add_run(f"Risk Score: {data.get('risk_score', 0)} / 100")
    score_run.bold = True
    score_run.font.size = Pt(18)

    grade = data.get("risk_grade", "Unknown")
    grade_para = doc.add_paragraph()
    grade_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grade_run = grade_para.add_run(grade)
    grade_run.bold = True
    grade_run.font.size = Pt(14)
    if grade in GRADE_COLORS and GRADE_COLORS[grade]:
        grade_run.font.color.rgb = GRADE_COLORS[grade]

    doc.add_page_break()


def _add_file_info(doc: Document, data: Dict):
    """Add file information and code statistics."""
    doc.add_heading("File Information", level=2)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"

    cells = [
        ("File Name", data.get("file", "N/A")),
        ("Language", data.get("language", "N/A").upper()),
        ("Total Findings", str(len(data.get("findings", [])))),
        ("Job ID", data.get("job_id", "N/A")),
    ]

    for i, (label, value) in enumerate(cells):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_paragraph("")


def _add_risk_score(doc: Document, data: Dict):
    """Add risk score analysis section."""
    doc.add_heading("Risk Score Analysis", level=2)

    formula = data.get("deduction_formula", "")
    if formula:
        doc.add_paragraph(f"Formula: {formula}")

    severity_counts = data.get("severity_counts", {})
    if severity_counts:
        doc.add_paragraph(
            f"Critical: {severity_counts.get('critical', 0)} | "
            f"High: {severity_counts.get('high', 0)} | "
            f"Medium: {severity_counts.get('medium', 0)} | "
            f"Low: {severity_counts.get('low', 0)}"
        )

    doc.add_paragraph("")


def _add_owasp_summary(doc: Document, findings: List[Dict]):
    """Add OWASP Top 10 summary table."""
    doc.add_heading("OWASP Top 10 Summary", level=2)

    # Group by OWASP category
    groups = {}
    for f in findings:
        cat = f.get("owasp") or "Unmapped"
        if cat not in groups:
            groups[cat] = {"total": 0, "severities": {}}
        groups[cat]["total"] += 1
        sev = (f.get("ai_severity") or f.get("severity") or "low").lower()
        groups[cat]["severities"][sev] = groups[cat]["severities"].get(sev, 0) + 1

    if not groups:
        doc.add_paragraph("No OWASP-mapped findings detected.")
        return

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "OWASP Category"
    headers[1].text = "Count"
    headers[2].text = "Severity Breakdown"

    for cat, data in sorted(groups.items(), key=lambda x: -x[1]["total"]):
        row = table.add_row().cells
        row[0].text = cat
        row[1].text = str(data["total"])
        row[2].text = ", ".join(f"{k}: {v}" for k, v in data["severities"].items())

    doc.add_paragraph("")


def _add_detailed_findings(doc: Document, findings: List[Dict]):
    """Add detailed findings section with code snippets and AI analysis."""
    doc.add_heading("Detailed Findings", level=2)

    if not findings:
        doc.add_paragraph("No findings to report.")
        return

    for i, f in enumerate(findings, 1):
        sev = (f.get("ai_severity") or f.get("severity") or "low").upper()
        tool = f.get("tool", "unknown")
        rule = f.get("rule_id", "N/A")

        heading = doc.add_heading(f"Finding #{i}: {rule} [{sev}]", level=3)

        # Basic info
        doc.add_paragraph(f"Tool: {tool} | Line: {f.get('line', 'N/A')} | CWE: {f.get('cwe', 'N/A')}")
        doc.add_paragraph(f"Message: {f.get('message', '')}")

        # Code snippet
        snippet = f.get("code_snippet")
        if snippet:
            doc.add_paragraph("Code:", style="Intense Quote")
            code_para = doc.add_paragraph()
            code_run = code_para.add_run(snippet)
            code_run.font.name = "Consolas"
            code_run.font.size = Pt(9)

        # AI analysis
        if f.get("ai_verdict"):
            doc.add_paragraph(f"AI Verdict: {f['ai_verdict']}")
        if f.get("ai_explanation"):
            doc.add_paragraph(f"Explanation: {f['ai_explanation']}")
        if f.get("ai_remediation"):
            doc.add_paragraph(f"Recommended Fix: {f['ai_remediation']}")
        if f.get("owasp"):
            doc.add_paragraph(f"OWASP Category: {f['owasp']}")

        doc.add_paragraph("")  # Spacer


# === Entry Point ===
if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python reportGenerator.py <input.json> <output.docx>", file=sys.stderr)
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    with open(input_path, "r") as f:
        report_data = json.load(f)

    result_path = generate_report(report_data, output_path)
    print(json.dumps({"status": "success", "path": result_path}))
